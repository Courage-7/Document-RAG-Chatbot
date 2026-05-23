from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

# initialize splitter (you can adjust later)
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1000,
    chunk_overlap=200
)

def docx_loader(file_path):
    doc = Document(file_path)

    content = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            content.append(para.text.strip())

    # Extract tables (important for RAG)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            content.append(" | ".join(row_text))

    return "\n".join(content)


def load_and_chunk_docx(file_path):
    # Step 1: extract text
    text = docx_loader(file_path)

    # Step 2: split into chunks
    chunks = text_splitter.split_text(text)

    return chunks